#!/usr/bin/env python3
"""
신학과사회 2025년 형식 DOCX 생성기
python-docx를 사용하여 정확한 형식의 Word 문서 생성
"""

import json
import sys
import os
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Mm, Inches, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print(json.dumps({
        "error": "python-docx not installed",
        "fix": "pip install python-docx"
    }))
    sys.exit(1)


# 2025년 신학과사회 형식 설정
SHINSA_2025 = {
    "page": {
        "width_mm": 152,    # 신국판
        "height_mm": 225,
        "margin_top_mm": 24,
        "margin_bottom_mm": 25,
        "margin_left_mm": 25,
        "margin_right_mm": 23,
    },
    "fonts": {
        "korean": "바탕",
        "english": "Times New Roman",
        "title_size": 14,
        "subtitle_size": 12.3,
        "author_size": 11,
        "body_size": 10.3,
        "abstract_size": 8.5,
        "abstract_title_size": 9,
        "section_title_size": 13,
        "footnote_size": 8.5,
        "header_size": 8.1,
    },
    "line_spacing": 1.6,  # 160%
    "journal_name": "신학과 사회",
}


def set_korean_font(run, font_name="바탕", size_pt=10.3):
    """한글 폰트 설정"""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    # 한글 폰트 설정 (eastAsia)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)


def add_paragraph_spacing(paragraph, line_spacing=1.6, space_after_pt=6):
    """줄간격 및 문단 간격 설정"""
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.space_after = Pt(space_after_pt)


def create_shinsa_docx(data: dict, output_path: str) -> dict:
    """
    신학과사회 형식 DOCX 생성

    Args:
        data: {
            title, subtitle, author, affiliation, field, email, funding,
            abstract_kr, keywords_kr, body, references,
            abstract_en, keywords_en,
            volume, issue, year, start_page, end_page
        }
        output_path: 저장 경로

    Returns:
        {success, path, message}
    """

    doc = Document()
    cfg = SHINSA_2025

    # ===== 페이지 설정 (신국판) =====
    section = doc.sections[0]
    section.page_width = Mm(cfg["page"]["width_mm"])
    section.page_height = Mm(cfg["page"]["height_mm"])
    section.top_margin = Mm(cfg["page"]["margin_top_mm"])
    section.bottom_margin = Mm(cfg["page"]["margin_bottom_mm"])
    section.left_margin = Mm(cfg["page"]["margin_left_mm"])
    section.right_margin = Mm(cfg["page"]["margin_right_mm"])

    # ===== 헤더 설정 =====
    volume = data.get("volume", 39)
    issue = data.get("issue", 2)
    year = data.get("year", 2025)
    start_page = data.get("start_page", 1)
    end_page = data.get("end_page", start_page + 20)

    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_run = header_para.add_run(f"「{cfg['journal_name']}」 {volume}({issue}) {year}")
    set_korean_font(header_run, cfg["fonts"]["korean"], cfg["fonts"]["header_size"])

    # 페이지 범위 (두 번째 줄)
    page_range_para = header.add_paragraph()
    page_range_run = page_range_para.add_run(f"pp. {start_page} - {end_page}")
    set_korean_font(page_range_run, cfg["fonts"]["korean"], cfg["fonts"]["header_size"])
    page_range_para.paragraph_format.space_after = Pt(12)

    # ===== 제목 (14pt, 가운데 정렬, 굵게) =====
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(data.get("title", ""))
    set_korean_font(title_run, cfg["fonts"]["korean"], cfg["fonts"]["title_size"])
    title_run.bold = True
    title_para.paragraph_format.space_before = Pt(20)
    title_para.paragraph_format.space_after = Pt(10)

    # 부제 (있는 경우)
    if data.get("subtitle"):
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.add_run(data["subtitle"])
        set_korean_font(subtitle_run, cfg["fonts"]["korean"], cfg["fonts"]["subtitle_size"])
        subtitle_para.paragraph_format.space_after = Pt(15)

    # ===== 저자 (11pt, 가운데, 띄어쓰기) =====
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 저자명 띄어쓰기 (이민규 → 이 민 규)
    author_name = data.get("author", "")
    author_spaced = " ".join(author_name) if author_name else ""
    author_run = author_para.add_run(author_spaced)
    set_korean_font(author_run, cfg["fonts"]["korean"], cfg["fonts"]["author_size"])
    author_run.bold = True

    # 저자 각주 기호
    footnote_symbol = "**" if data.get("funding") else "*"
    symbol_run = author_para.add_run(footnote_symbol)
    symbol_run.font.superscript = True
    set_korean_font(symbol_run, cfg["fonts"]["korean"], 8)

    author_para.paragraph_format.space_after = Pt(20)

    # ===== 구분선 =====
    doc.add_paragraph("─" * 40)

    # ===== 국문초록 (붙여쓰기) =====
    abstract_title_para = doc.add_paragraph()
    abstract_title_run = abstract_title_para.add_run("국문초록")
    set_korean_font(abstract_title_run, cfg["fonts"]["korean"], cfg["fonts"]["abstract_title_size"])
    abstract_title_run.bold = True
    abstract_title_para.paragraph_format.space_after = Pt(8)

    # 초록 본문 (8.5pt)
    abstract_para = doc.add_paragraph()
    abstract_run = abstract_para.add_run(data.get("abstract_kr", "[초록 작성 필요]"))
    set_korean_font(abstract_run, cfg["fonts"]["korean"], cfg["fonts"]["abstract_size"])
    add_paragraph_spacing(abstract_para, 1.5, 8)
    abstract_para.paragraph_format.left_indent = Mm(10)
    abstract_para.paragraph_format.right_indent = Mm(10)

    # 주제어
    keywords_para = doc.add_paragraph()
    keywords_label = keywords_para.add_run("주제어: ")
    set_korean_font(keywords_label, cfg["fonts"]["korean"], cfg["fonts"]["abstract_size"])
    keywords_label.bold = True

    keywords_kr = data.get("keywords_kr", [])
    keywords_text = ", ".join(keywords_kr) if keywords_kr else "[주제어 5개]"
    keywords_run = keywords_para.add_run(keywords_text)
    set_korean_font(keywords_run, cfg["fonts"]["korean"], cfg["fonts"]["abstract_size"])
    keywords_para.paragraph_format.left_indent = Mm(10)

    # ===== 구분선 =====
    doc.add_paragraph("─" * 40)

    # ===== 본문 =====
    body_text = data.get("body", "")
    sections = parse_body_sections(body_text)

    for sec in sections:
        # 섹션 제목
        if sec["title"]:
            sec_para = doc.add_paragraph()
            sec_run = sec_para.add_run(f"{sec['number']} {sec['title']}")

            if sec["level"] == 1:
                # 장 제목 (Ⅰ. 서론) - 13pt
                set_korean_font(sec_run, cfg["fonts"]["korean"], cfg["fonts"]["section_title_size"])
                sec_run.bold = True
                sec_para.paragraph_format.space_before = Pt(18)
            elif sec["level"] == 2:
                # 절 제목 (1. 절제목) - 11pt
                set_korean_font(sec_run, cfg["fonts"]["korean"], 11)
                sec_run.bold = True
                sec_para.paragraph_format.space_before = Pt(12)
            else:
                # 항 제목 (1) 소제목) - 10.3pt
                set_korean_font(sec_run, cfg["fonts"]["korean"], cfg["fonts"]["body_size"])
                sec_para.paragraph_format.space_before = Pt(8)

            sec_para.paragraph_format.space_after = Pt(6)

        # 본문 내용
        if sec["content"]:
            for para_text in sec["content"].split("\n\n"):
                if para_text.strip():
                    body_para = doc.add_paragraph()
                    body_run = body_para.add_run(para_text.strip())
                    set_korean_font(body_run, cfg["fonts"]["korean"], cfg["fonts"]["body_size"])
                    add_paragraph_spacing(body_para, cfg["line_spacing"], 6)
                    body_para.paragraph_format.first_line_indent = Mm(5)

    # ===== 구분선 =====
    doc.add_paragraph("─" * 40)

    # ===== 참고문헌 =====
    ref_title = doc.add_paragraph()
    ref_title_run = ref_title.add_run("참고문헌")
    set_korean_font(ref_title_run, cfg["fonts"]["korean"], cfg["fonts"]["section_title_size"])
    ref_title_run.bold = True
    ref_title.paragraph_format.space_before = Pt(18)
    ref_title.paragraph_format.space_after = Pt(10)

    references = data.get("references", [])

    # 국문/외국어 분류
    korean_refs = [r for r in references if r and r[0] >= '가' and r[0] <= '힣']
    foreign_refs = [r for r in references if r and not (r[0] >= '가' and r[0] <= '힣')]

    if korean_refs:
        kr_header = doc.add_paragraph()
        kr_header_run = kr_header.add_run("<국문 자료>")
        set_korean_font(kr_header_run, cfg["fonts"]["korean"], cfg["fonts"]["body_size"])
        kr_header_run.bold = True

        for ref in korean_refs:
            ref_para = doc.add_paragraph()
            ref_run = ref_para.add_run(ref)
            set_korean_font(ref_run, cfg["fonts"]["korean"], cfg["fonts"]["body_size"])
            ref_para.paragraph_format.left_indent = Mm(5)
            ref_para.paragraph_format.first_line_indent = Mm(-5)

    if foreign_refs:
        en_header = doc.add_paragraph()
        en_header_run = en_header.add_run("<외국어 자료>")
        set_korean_font(en_header_run, cfg["fonts"]["korean"], cfg["fonts"]["body_size"])
        en_header_run.bold = True
        en_header.paragraph_format.space_before = Pt(10)

        for ref in foreign_refs:
            ref_para = doc.add_paragraph()
            ref_run = ref_para.add_run(ref)
            set_korean_font(ref_run, cfg["fonts"]["english"], cfg["fonts"]["body_size"])
            ref_para.paragraph_format.left_indent = Mm(5)
            ref_para.paragraph_format.first_line_indent = Mm(-5)

    # ===== 구분선 =====
    doc.add_paragraph("─" * 40)

    # ===== 영문 초록 =====
    abstract_en_title = doc.add_paragraph()
    abstract_en_title_run = abstract_en_title.add_run("Abstract")
    set_korean_font(abstract_en_title_run, cfg["fonts"]["english"], cfg["fonts"]["abstract_title_size"])
    abstract_en_title_run.bold = True

    abstract_en_para = doc.add_paragraph()
    abstract_en_run = abstract_en_para.add_run(data.get("abstract_en", "[Abstract required]"))
    set_korean_font(abstract_en_run, cfg["fonts"]["english"], cfg["fonts"]["abstract_size"])
    add_paragraph_spacing(abstract_en_para, 1.5, 8)
    abstract_en_para.paragraph_format.left_indent = Mm(10)
    abstract_en_para.paragraph_format.right_indent = Mm(10)

    # Keywords
    keywords_en_para = doc.add_paragraph()
    keywords_en_label = keywords_en_para.add_run("Keywords: ")
    set_korean_font(keywords_en_label, cfg["fonts"]["english"], cfg["fonts"]["abstract_size"])
    keywords_en_label.bold = True

    keywords_en = data.get("keywords_en", [])
    keywords_en_text = ", ".join(keywords_en) if keywords_en else "[5 keywords]"
    keywords_en_run = keywords_en_para.add_run(keywords_en_text)
    set_korean_font(keywords_en_run, cfg["fonts"]["english"], cfg["fonts"]["abstract_size"])
    keywords_en_para.paragraph_format.left_indent = Mm(10)

    # ===== 저자 각주 =====
    doc.add_paragraph("─" * 40)

    footnote_para = doc.add_paragraph()

    # 연구비 지원 (있는 경우)
    if data.get("funding"):
        funding_run = footnote_para.add_run(f"* {data['funding']}\n")
        set_korean_font(funding_run, cfg["fonts"]["korean"], cfg["fonts"]["footnote_size"])

    # 저자 정보
    author_info_parts = [data.get("affiliation", "")]
    if data.get("field"):
        author_info_parts.append(data["field"])
    if data.get("email"):
        author_info_parts.append(data["email"])

    author_info = "/ ".join(filter(None, author_info_parts))
    author_footnote_run = footnote_para.add_run(f"{footnote_symbol} {author_info}")
    set_korean_font(author_footnote_run, cfg["fonts"]["korean"], cfg["fonts"]["footnote_size"])

    # ===== 저장 =====
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    return {
        "success": True,
        "path": str(output_path.absolute()),
        "message": f"신학과사회 2025년 형식 DOCX 생성 완료: {output_path.name}"
    }


def parse_body_sections(body: str) -> list:
    """본문에서 섹션 구조 추출"""
    sections = []
    lines = body.split("\n")

    current = {"level": 0, "number": "", "title": "", "content": []}
    chapter_num = 0
    section_num = 0
    subsection_num = 0

    roman_numerals = ["", "Ⅰ", "Ⅱ", "Ⅲ", "Ⅳ", "Ⅴ", "Ⅵ", "Ⅶ", "Ⅷ", "Ⅸ", "Ⅹ"]

    import re

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # # 또는 I. II. 패턴 감지
        h1_match = re.match(r'^#\s+(.+)$', line)
        roman_match = re.match(r'^(I{1,3}|IV|V|VI{0,3})\.\s+(.+)$', line)
        h2_match = re.match(r'^##\s+(.+)$', line)
        arabic_match = re.match(r'^(\d+)\.\s+(.+)$', line)
        h3_match = re.match(r'^###\s+(.+)$', line)
        paren_match = re.match(r'^(\d+)\)\s+(.+)$', line)

        if h1_match or roman_match:
            if current["title"] or current["content"]:
                current["content"] = "\n".join(current["content"])
                sections.append(current)

            chapter_num += 1
            section_num = 0
            subsection_num = 0
            title = h1_match.group(1) if h1_match else roman_match.group(2)
            current = {
                "level": 1,
                "number": f"{roman_numerals[chapter_num]}.",
                "title": title,
                "content": []
            }
        elif h2_match or (arabic_match and not line.startswith('1)')):
            if current["title"] or current["content"]:
                current["content"] = "\n".join(current["content"])
                sections.append(current)

            section_num += 1
            subsection_num = 0
            title = h2_match.group(1) if h2_match else arabic_match.group(2)
            current = {
                "level": 2,
                "number": f"{section_num}.",
                "title": title,
                "content": []
            }
        elif h3_match or paren_match:
            if current["title"] or current["content"]:
                current["content"] = "\n".join(current["content"])
                sections.append(current)

            subsection_num += 1
            title = h3_match.group(1) if h3_match else paren_match.group(2)
            current = {
                "level": 3,
                "number": f"{subsection_num})",
                "title": title,
                "content": []
            }
        else:
            current["content"].append(line)

    # 마지막 섹션 추가
    if current["title"] or current["content"]:
        current["content"] = "\n".join(current["content"])
        sections.append(current)

    return sections


if __name__ == "__main__":
    # 명령줄에서 JSON 입력 받기
    if len(sys.argv) < 3:
        print(json.dumps({
            "error": "Usage: python create_docx.py <input.json> <output.docx>"
        }))
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]

    try:
        with open(input_file, 'r', encoding='utf-8') as f:
            data = json.load(f)

        result = create_shinsa_docx(data, output_file)
        print(json.dumps(result, ensure_ascii=False))

    except Exception as e:
        print(json.dumps({
            "error": str(e)
        }, ensure_ascii=False))
        sys.exit(1)

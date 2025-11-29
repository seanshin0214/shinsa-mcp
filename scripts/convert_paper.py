#!/usr/bin/env python3
"""
신관수 GPTs 리더십 논문 → 신학과사회 2025년 형식 DOCX 변환
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path

def set_korean_font(run, font_name="바탕", size_pt=10.3):
    """한글 폰트 설정"""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)

def add_paragraph_spacing(paragraph, line_spacing=1.6, space_after_pt=6):
    """줄간격 및 문단 간격 설정"""
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.space_after = Pt(space_after_pt)

# 논문 데이터
paper_data = {
    "title": "GPTs 기반 AI 리더십 의사결정 도구의 실무 활용에 관한 질적 연구",
    "subtitle": "300권 리더십 문헌 분석과 143명의 AI 전문가 시스템 기반 도구 개발",
    "author": "신관수",
    "affiliation": "한일장신대학교 전북소피아국제대학",
    "field": "리더십학",
    "email": "",

    "abstract_kr": """본 연구는 GPTs(Generative Pre-trained Transformers) 기반 AI 리더십 의사결정 도구의 개발과 실무 활용 경험을 탐구하였다. 300권의 리더십 문헌을 체계적으로 분석하여 30,000개의 리더십 도전 과제를 도출하고, 이를 15개 핵심 카테고리로 분류하였다. 연구진은 RAG(Retrieval-Augmented Generation)와 MCP(Model Context Protocol)를 결합한 143명의 지식기반 AI 전문가 시스템을 구축한 후, 이를 일반 사용자가 접근 가능한 20개의 GPTs로 축약하여 개발하였다. 보안 분야 전문가 2명(신한금융그룹 보안담당, 대기업 보안회사 대표), 삼성전자 플랜트 개발 부문 진동분야 수석연구원 1명, 디자인 회사 대표 1명, 요양보호센터 대표 1명 등 총 5명의 현업 리더를 대상으로 1주일간의 실무 활용 후 심층 인터뷰를 실시하여 해석적 현상학적 분석(IPA)을 수행하였다. 연구 결과, 참여자들은 GPTs 도구가 의사결정의 다각적 관점 확보, 인지적 맹점 식별, 의사결정 품질 향상에 기여한다고 인식하였다. 동시에 산업별 맥락 이해의 한계, 인간적 판단의 필수성, 기술 의존성에 대한 우려도 제기되었다. 본 연구는 GPTs를 리더십 의사결정에 적용한 최초의 실증 연구로서, 증강 지능 이론의 확장과 AI-인간 협업 리더십 모델 개발에 기여한다.""",

    "keywords_kr": ["GPTs", "AI 증강 리더십", "의사결정", "질적 연구", "해석적 현상학적 분석"],

    "abstract_en": """This study explores the development and practical implementation experiences of a GPTs-based AI leadership decision-making tool. Through systematic analysis of 300 leadership texts, 30,000 leadership challenge problems were derived and categorized into 15 core domains. The research team first built 143 knowledge-based AI expert personas combining RAG (Retrieval-Augmented Generation) and MCP (Model Context Protocol), then condensed them into 20 accessible GPTs for general users. Interpretative Phenomenological Analysis (IPA) was conducted through in-depth interviews with five industry leaders after one week of practical use: two security executives (Shinhan Financial Group security director and a major enterprise security company CEO), one senior researcher in vibration analysis at Samsung Electronics Plant Development Division, one design company CEO, and one care center CEO. Findings reveal that participants perceived the GPTs tool as contributing to multi-perspective consideration, identification of cognitive blind spots, and improvement of decision quality. Concurrently, concerns were raised regarding limitations in industry-specific contextual understanding, the necessity of human judgment, and technology dependence. As the first empirical study applying GPTs to leadership decision-making, this research contributes to the extension of augmented intelligence theory and the development of AI-human collaborative leadership models.""",

    "keywords_en": ["GPTs", "AI-augmented leadership", "decision-making", "qualitative research", "Interpretative Phenomenological Analysis"],

    "volume": 39,
    "issue": 2,
    "year": 2025,
    "start_page": 1,
    "end_page": 30
}

# 본문 섹션 (마크다운에서 파싱)
sections = [
    {"level": 1, "number": "Ⅰ.", "title": "서론", "content": ""},
    {"level": 2, "number": "1.", "title": "연구 배경 및 문제 제기", "content": """인공지능 시대의 도래와 함께 리더십 의사결정의 환경은 근본적인 변화를 겪고 있다. 오늘날의 교육 및 기업 리더들은 정보 과부하, 변화의 가속화, 복잡성의 증가로 특징되는 전례 없는 도전에 직면하고 있다. 전통적인 리더십 이론들은 기초적이지만, 현대 AI 시스템이 제시하는 역량과 도전을 예측할 수 없는 시대에 개발되었다.

대규모 언어 모델(LLM)의 등장과 조직 워크플로우로의 통합은 인간 의사결정을 증강시키는 새로운 가능성을 창출하였다. 2023년 11월 OpenAI가 발표한 GPTs(Generative Pre-trained Transformers)는 사용자 맞춤형 AI 어시스턴트의 새로운 패러다임을 제시하였다. 이어 2024년 1월에 출시된 GPT Store는 3개월 만에 300만 개 이상의 커스텀 GPTs가 개발되는 폭발적인 성장을 보여주었다.

MIT의 Noy와 Zhang이 Science에 발표한 연구에 따르면, ChatGPT를 활용한 전문직 종사자들은 업무 완료 시간이 40% 단축되고 결과물의 품질이 18% 향상되었다. 그러나 이러한 기술의 리더십 맥락에서의 실제 적용에 관한 학술적 연구는 여전히 미개척 영역으로 남아 있다.

증강 지능(Augmented Intelligence)의 개념은 AI를 인간 지능의 대체물이 아닌 인지적 역량의 강화 도구로 이해하는 패러다임 전환을 대표한다. Herbert Simon의 제한된 합리성(bounded rationality)에 관한 기초 연구는 인간 의사결정자가 정보를 최적으로 처리하는 능력을 제약하는 인지적 한계 하에서 작동함을 확립하였다. 현대 조직 환경에서 이러한 한계는 리더들이 처리해야 하는 정보의 순수한 양과 속도에 의해 증폭된다."""},

    {"level": 2, "number": "2.", "title": "연구 목적", "content": """본 연구의 목적은 세 가지이다:

첫째, 문헌 분석 및 프레임워크 개발이다. 300권의 리더십 문헌을 체계적으로 분석하여 30,000개의 리더십 도전 과제를 도출하고 이를 기반으로 GPTs 개발을 위한 포괄적인 의사결정 프레임워크를 개발하는 것이다.

둘째, AI 전문가 시스템 구축이다. RAG와 MCP를 결합한 143명의 지식기반 AI 전문가 시스템을 구축하고, 이를 일반 사용자가 접근 가능한 20개의 GPTs로 축약하여 개발하는 것이다.

셋째, 실무 경험 탐구이다. 다양한 산업 분야의 현업 리더 5명을 대상으로 1주일간의 실무 활용 후 의사결정 과정에서의 체험 경험을 탐구하는 것이다."""},

    {"level": 2, "number": "3.", "title": "연구 문제", "content": """본 연구는 다음의 연구 문제에 의해 안내된다:

연구문제 1: 300권의 리더십 문헌에 대한 체계적 분석에서 어떤 의사결정 프레임워크와 패턴이 도출되는가?

연구문제 2: 다양한 산업의 리더들은 GPTs 기반 AI 의사결정 지원 도구를 그들의 리더십 실무에서 어떻게 경험하는가?

연구문제 3: GPTs 증강 리더십 의사결정의 인식된 가능성과 한계는 무엇인가?"""},

    {"level": 1, "number": "Ⅱ.", "title": "이론적 배경", "content": ""},
    {"level": 2, "number": "1.", "title": "리더십 의사결정 이론", "content": """리더십 의사결정의 이론적 기초는 다양한 학문적 전통에서 유래한다. Burns와 Bass가 발전시킨 변혁적 리더십 이론은 비전적 의사결정을 통해 추종자들에게 영감을 주고 조직 변화를 창출하는 리더의 역할을 강조한다. 이 관점은 가치 기반 의사결정의 중요성과 그들의 선택이 조직 문화와 이해관계자 복지에 미치는 광범위한 영향을 고려해야 하는 리더의 책임을 강조한다.

Greenleaf가 제시한 서번트 리더십은 추종자들의 필요를 우선시하고 타인에 대한 봉사에 기반한 윤리적 의사결정을 강조하는 대안적 프레임워크를 제공한다. Kouzes와 Posner의 연구는 모범적 리더십의 다섯 가지 실천을 제시하며 리더십 실행의 구체적 행동을 명시하였다.

Kahneman의 이중 과정 인지 이론은 빠르고 직관적인 시스템 1 사고와 느리고 숙고적인 시스템 2 처리를 구분하며, 현대적 의사결정 이해에 심대한 영향을 미쳤다. 이 프레임워크는 효과적인 의사결정이 인지적 편향에 대한 인식과 당면한 결정의 성격에 따라 직관적 및 분석적 과정 모두의 전략적 배치를 요구함을 시사한다."""},

    {"level": 2, "number": "2.", "title": "GPTs와 증강 지능", "content": """증강 지능의 개념은 AI를 인간 지능의 대체물로 보는 것에서 인간 인지 역량을 강화하는 도구로 이해하는 패러다임 전환을 대표한다. Davenport와 Kirby는 조직적 맥락에서 AI의 가장 가치 있는 적용은 인간 판단을 대체하기보다 보완하는 것임을 주장하였다.

GPTs는 2023년 11월 OpenAI가 발표한 사용자 맞춤형 ChatGPT의 새로운 형태로, 코딩 경험 없이도 특정 목적에 맞는 AI 어시스턴트를 구축할 수 있게 한다. GPT Store의 출시 이후 사용자들은 DALL·E, 글쓰기, 연구, 프로그래밍, 교육, 라이프스타일 등 다양한 범주의 300만 개 이상의 커스텀 GPTs를 개발하였다.

기술수용모델(TAM)은 사용자가 새로운 기술을 채택하는 방식을 이해하기 위한 이론적 프레임워크를 제공하며, 지각된 유용성과 지각된 사용 용이성을 기술 수용의 핵심 결정요인으로 강조한다. 이후 확장된 통합기술수용이론(UTAUT)은 사회적 영향, 촉진 조건, 개인적 특성과 같은 추가적 요인을 통합하였다."""},

    {"level": 2, "number": "3.", "title": "연구 공백 및 기여", "content": """AI의 리더십 적용에 대한 관심이 증가하고 있음에도 불구하고, 문헌에는 여전히 몇 가지 중요한 공백이 남아 있다. 첫째, 대규모 리더십 문헌 코퍼스를 체계적으로 분석하여 AI 구현 가능한 의사결정 프레임워크를 도출한 기존 연구가 없다. 둘째, GPTs 기술 연구가 등장하였으나 리더십 맥락에서의 적용을 검토한 연구는 없다. 셋째, AI 의사결정 지원 도구에 대한 리더들의 체험 경험을 탐구하는 질적 연구는 문헌에서 현저하게 부재하다."""},

    {"level": 1, "number": "Ⅲ.", "title": "연구 방법", "content": ""},
    {"level": 2, "number": "1.", "title": "연구 설계", "content": """본 연구는 문헌 분석, 도구 개발, 그리고 해석적 현상학적 분석(IPA)을 결합한 다단계 연구 설계를 채택하였다. 연구는 네 가지 단계로 진행되었다: (1) 체계적 문헌 분석, (2) AI 도구 개발, (3) 도구 실행, (4) 현상학적 탐구.

해석적 현상학적 분석(IPA)은 체험 경험의 이해와 개인이 경험하는 현상의 본질적 구조를 밝히는 데 강조점을 두기 때문에 질적 연구 요소를 위해 선택되었다. 이 접근법은 리더들이 AI 의사결정 지원 도구와의 상호작용을 어떻게 경험하고 의미를 부여하는지를 조사하는 데 특히 적합하다."""},

    {"level": 2, "number": "2.", "title": "문헌 분석 단계", "content": """문헌 분석 단계는 목적적 표집을 통해 선정된 300권의 리더십 문헌에 대한 체계적 검토를 포함하였다. 선정 기준은 인정받는 리더십 학자 또는 실무자에 의한 출판, 의사결정 과정에 대한 실질적 다룸, 컴퓨터 분석을 위한 디지털 형식의 가용성, 기초 저작을 포함하면서 현대적 관련성을 보장하기 위해 지난 40년 내 출판이었다.

코퍼스에는 Burns, Bass, Greenleaf, Kouzes와 Posner, Collins, Senge, Kotter 및 다른 영향력 있는 리더십 저자들의 대표 저작이 포함되었다. 분석은 자연어 처리 기법을 활용하여 코퍼스 전반에 걸친 반복 주제, 프레임워크, 의사결정 패턴을 식별하였다. 그 결과 30,000개의 리더십 도전 과제가 도출되어 15개의 핵심 카테고리로 분류되었다."""},

    {"level": 2, "number": "3.", "title": "GPTs 도구 개발 단계", "content": """GPTs 기반 AI 도구는 두 단계로 개발되었다. 1단계에서는 RAG(Retrieval-Augmented Generation)와 MCP(Model Context Protocol)를 결합하여 143명의 월드클래스 수준 지식기반 AI 전문가 시스템을 구축하였다. 각 전문가 페르소나는 특정 리더십 영역의 지식, 기술, 경험을 구현하도록 설계되었다.

2단계에서는 143명의 AI 전문가 시스템을 일반 사용자가 쉽게 접근하고 활용할 수 있도록 20개의 GPTs로 축약하였다. 이 과정에서 핵심 지식과 기술을 문서화하고, 명확한 지침을 개발하여 사용자 친화적인 인터페이스를 구현하였다. 도구는 문헌 분석에서 도출된 의사결정 프레임워크를 운용화하면서 각 결정의 특정 맥락에 반응하는 구조화된 의사결정 과정을 통해 리더를 안내하도록 설계되었다."""},

    {"level": 2, "number": "4.", "title": "연구 참여자", "content": """참여자는 다양한 산업 맥락의 경험 있는 리더를 포함하기 위해 목적적 표집을 통해 모집되었다. 선정 기준은 최소 10년의 리더십 경험, 중요한 의사결정 책임을 수반하는 현재 역할, 기본적인 기술적 문해력, 1주일의 기간 동안 GPTs 도구에 참여할 의향이었다.

최종 표본은 5명으로 구성되었다: 신한금융그룹 보안담당 임원, 대기업 보안전문회사 대표, 삼성전자 플랜트개발부문 진동분야 수석연구원, 디자인 회사 대표, 요양보호센터 대표. 참여자들의 연령은 42-58세이며, 리더십 경험은 10-25년이었다."""},

    {"level": 2, "number": "5.", "title": "자료 수집", "content": """자료 수집은 세 가지 방법을 채택하였다. 첫째, 심층 인터뷰로서 반구조화된 인터뷰는 1주일의 실행 기간 종료 시점에 실시되었으며, 각각 60-90분이 소요되었다. 인터뷰 프로토콜은 현상학적 면담 원칙에 기반하여 개발되었으며, 참여자들의 전반적인 경험에 관한 광범위한 질문으로 시작하여 GPTs 증강 의사결정의 특정 측면에 점진적으로 집중하였다.

둘째, 도구 사용 로그를 통해 도구 사용의 빈도와 패턴에 관한 양적 데이터를 포착하였다. 셋째, 성찰 일지를 통해 실행 기간 전반에 걸쳐 참여자들의 경험과 성찰에 대한 지속적인 문서화를 제공하였다.

모든 인터뷰는 오디오 녹음되어 축어 전사되었다. 참여자들은 사전 동의서를 제공하였으며, 연구는 기관생명윤리위원회의 승인을 받았다."""},

    {"level": 2, "number": "6.", "title": "자료 분석", "content": """인터뷰 전사본과 일지 기록은 Braun과 Clarke가 제시한 6단계 접근법에 따라 주제 분석을 수행하였다: (1) 데이터 친숙화, (2) 초기 코드 생성, (3) 주제 탐색, (4) 주제 검토, (5) 주제 정의 및 명명, (6) 보고서 작성. 분석은 반복적으로 수행되었으며, 참여자 간 및 산업 분야별 하위 그룹 간 지속적 비교가 이루어졌다.

신뢰성 확보 전략으로 참여자 확인(member checking), 동료 검토(peer debriefing), 상세한 감사 추적(audit trail) 유지가 사용되었다. 두 명의 연구자가 전사본 하위 집합을 독립적으로 코딩하여 87%의 평가자 간 일치도를 달성하였다. 불일치는 토론과 합의를 통해 해결되었다."""},

    {"level": 1, "number": "Ⅳ.", "title": "연구 결과", "content": ""},
    {"level": 2, "number": "1.", "title": "문헌에서 도출된 리더십 의사결정 프레임워크", "content": """300권의 리더십 문헌 분석은 30,000개의 리더십 도전 과제를 도출하였으며, 이는 15개의 상호 연결된 차원으로 구성된 포괄적인 의사결정 프레임워크로 분류되었다.

핵심 차원은 다음과 같다. 이해관계자 고려는 결정에 영향을 받는 모든 당사자의 이해를 식별하고 비중을 두는 것의 중요성을 강조한다. 윤리적 평가는 서번트 리더십과 가치 기반 의사결정 전통의 원칙을 활용한다. 전략적 정렬은 결정이 더 넓은 조직 미션과 목표에 연결되도록 보장한다. 실행 가능성은 실제적 제약과 자원 요구사항을 다룬다. 결과 예측은 잠재적 결정의 의도된 결과와 의도하지 않은 결과 모두를 예상한다.

세 가지 메타 패턴이 도출되었다: 다양한 관점을 동시에 유지하는 역량인 통합적 사고, 단기 및 장기 함의의 균형인 시간적 고려, 결정의 지속적 평가와 조정인 성찰적 실천. 이러한 패턴은 구조화된 질문 프로토콜과 분석적 프레임워크를 통해 GPTs 도구에서 운용화되었다."""},

    {"level": 2, "number": "2.", "title": "GPTs 기반 AI 도구에 대한 리더들의 경험", "content": """인터뷰 데이터에 대한 주제 분석은 참여자들의 GPTs 의사결정 지원 도구 경험에 관한 세 가지 주요 주제를 밝혔다.

첫째, 향상된 관점 취하기이다. 참여자들은 일관되게 GPTs 도구가 그렇지 않았다면 간과했을 수 있는 관점에 대한 고려를 촉진했다고 보고하였다. 신한금융그룹 보안담당 임원은 "이 도구는 보안 정책 변경 결정을 내릴 때 처음에 고려하지 않았던 이해관계자들에 대해 생각하도록 계속 밀어붙였습니다. IT 부서와 경영진뿐만 아니라 현장 직원들과 외부 파트너들에 대한 영향까지 고려하게 되었습니다"라고 언급하였다.

둘째, 인지적 맹점 식별이다. 도구가 참여자들의 의사결정에서 인지적 맹점과 편향을 밝히는 역량과 관련되었다. 대기업 보안전문회사 대표는 "비용 절감 조치에 대해 결정이 단순명료하다고 확신했습니다. 그러나 도구가 장기적인 보안 인력 역량과 승계 계획에 대해 질문하기 시작했을 때, 당장의 비용 필요에만 전적으로 집중하고 있었음을 깨달았습니다"라고 보고하였다.

셋째, 보존된 주체성이다. 도구가 그들의 사고에 미치는 영향에도 불구하고, 참여자들은 그들의 결정에 대한 완전한 주체성을 유지했다고 일관되게 강조하였다. GPTs는 의사결정자가 아닌 사고 파트너로 경험되었다."""},

    {"level": 2, "number": "3.", "title": "인식된 도전과 한계", "content": """참여자들이 전반적으로 긍정적인 경험을 보고한 반면, GPTs 증강 의사결정에서 중요한 도전과 한계도 식별하였다.

맥락적 한계와 관련하여, 참여자들은 GPTs 도구가 그들의 특정 산업 맥락의 중요한 뉘앙스를 파악하지 못하는 경우를 언급하였다. 신한금융그룹 보안담당 임원은 "도구가 이해관계자 커뮤니케이션에 관해 매우 일반적인 질문을 주었지만, 우리 금융산업의 특수한 규제 환경이나 특정 부서 간의 갈등 역사를 이해하지 못했습니다"라고 언급하였다.

인간 판단의 환원 불가능성도 언급되었다. 특히 대인 관계 역학이나 조직 문화를 포함하는 결정에서 자주 강조되었다. 디자인 회사 대표는 "어떤 결정은 궁극적으로 사람에 대한 직감에 달려 있습니다. GPTs는 그것에 도움이 될 수 없었습니다"라고 언급하였다.

기술 의존에 대한 우려도 참여자의 약 절반에서 제기되었다. 삼성전자 수석연구원은 "이 도구를 정기적으로 사용하면 스스로 철저한 분석을 하는 것에 게을러질 수 있을지 궁금합니다"라고 우려를 표명하였다."""},

    {"level": 2, "number": "4.", "title": "산업별 비교 분석", "content": """산업별 하위 그룹 간 비교는 경험에서 유사점과 주목할 만한 차이점 모두를 드러내었다. 모든 그룹은 도구의 관점 취하기 향상과 맹점 식별 역량을 가치 있게 평가하였다.

차이점으로는 보안 분야 리더들이 규제 준수와 위험 관리의 윤리적 차원에 더 큰 강조점을 두었고, 기술 분야 리더들은 혁신과 효율성 함의에 더 자주 집중하였다.

도구 사용 패턴에서 보안 분야는 결정당 평균 26.3분의 상호작용 시간으로 보다 숙고적인 참여 스타일을 보였고, 기타 분야는 결정당 18.7분으로 확장된 숙고보다 빠른 상담을 선호하였다."""},

    {"level": 1, "number": "Ⅴ.", "title": "논의", "content": ""},
    {"level": 2, "number": "1.", "title": "이론적 함의", "content": """본 연구의 결과는 리더십 맥락에서 AI-인간 협업적 의사결정의 실행 가능성에 대한 실증적 증거를 제공함으로써 증강 지능 이론을 확장한다. Davenport와 Kirby의 증강 지능 개념화는 인간과 기계 역량의 보완적 성격을 강조하였다. 본 연구는 인간 판단을 대체하기보다 강화하는 구조화된 GPTs 도구를 통해 이러한 보완성이 어떻게 운용화될 수 있는지를 보여준다.

보존된 주체성의 주제는 리더십 이론에 특히 의미 있다. 변혁적 리더십 이론은 개인적 자질과 판단이 조직 성공의 핵심인 비전적 변화 주도자로서의 리더 역할을 강조한다. 참여자들이 GPTs 지원으로부터 혜택을 받으면서도 완전한 의사결정 주체성을 유지했다는 발견은 AI 증강이 리더의 변혁적 역할을 감소시키지 않을 수 있으며 잠재적으로 리더십 실천에 이용 가능한 인지적 자원을 확장함으로써 이를 강화할 수 있음을 시사한다.

AI 상호작용을 통한 맹점 식별은 행동적 의사결정 이론에서 식별된 인지적 한계를 다루기 위한 새로운 메커니즘을 대표한다. Kahneman의 연구는 인지적 편향이 의사결정을 어떻게 체계적으로 왜곡하는지를 보여주었다. GPTs 도구의 소크라테스적 질문 접근법은 편향 제거 개입으로 기능하는 것으로 보인다."""},

    {"level": 2, "number": "2.", "title": "실무적 함의", "content": """실무자들에게 본 연구는 여러 실행 가능한 통찰을 제공한다. 첫째, GPTs 도구의 권고 엔진이 아닌 질문 파트너로서의 설계는 리더 참여와 주체성 유지에 핵심적인 것으로 보인다. GPTs 의사결정 지원을 구현하는 조직은 행동을 처방하기보다 성찰을 강화하는 설계를 고려해야 한다.

둘째, 식별된 맥락적 한계는 특정 산업 맥락에 맞게 GPTs 도구를 맞춤화하는 것의 중요성을 시사한다. 조직은 산업별 정보를 통합하는 GPTs 지식 기반 개발에 투자해야 한다.

셋째, 기술 의존에 대한 우려는 GPTs 증강 프로세스와 함께 리더들의 자율적 의사결정 역량을 유지하고 개발하는 의도적인 실천의 필요성을 시사한다. 조직은 기술 위축을 방지하기 위해 AI 지원 의사결정과 독립적 의사결정 사이를 번갈아가는 것을 고려할 수 있다."""},

    {"level": 2, "number": "3.", "title": "한계 및 향후 연구", "content": """본 연구는 결과 해석에 정보를 제공하고 향후 연구를 안내해야 할 몇 가지 한계를 가지고 있다. 표본 크기는 현상학적 탐구에 적합하나 일반화 가능성에 제한이 있으며, 대규모 표본 연구가 필요하다. 1주일의 실행 기간은 장기적 효과 포착에 불충분하며, 종단적 연구 설계가 필요하다. 자기 보고 의존은 사회적 바람직성 편향과 회고적 회상의 한계가 있으며, 관찰적 방법과 객관적 측정 통합이 필요하다. 한국 리더들로 제한된 문화적 맥락은 문화 간 비교 연구가 필요하다.

향후 연구는 리더십 개발과 의사결정 역량에 대한 지속적인 GPTs 도구 사용의 종단적 효과를 검토해야 한다. GPTs 증강 대 전통적 의사결정 프로세스의 결과를 비교하는 연구는 이러한 도구의 실제적 효과성에 대한 귀중한 증거를 제공할 것이다."""},

    {"level": 1, "number": "Ⅵ.", "title": "결론", "content": """본 연구는 리더십 의사결정 맥락에서 GPTs 적용에 대한 최초의 실증적 조사를 대표한다. 300권의 리더십 문헌에 대한 체계적 분석, RAG와 MCP 기반 143명의 AI 전문가 시스템 구축, 20개의 GPTs 도구 개발, 그리고 리더들의 경험에 대한 현상학적 탐구를 통해, 연구는 AI 증강 리더십의 신생 분야에 이론적 및 실무적 통찰 모두를 기여한다.

주요 발견으로는 질문 파트너로 설계된 GPTs 도구가 관점 취하기를 촉진하고, 인지적 맹점을 식별하며, 이해관계자와 결과에 대한 보다 체계적인 고려를 촉구함으로써 리더십 의사결정을 향상시킬 수 있다는 것이다. 이러한 혜택은 리더들의 주체성과 그들의 결정에 대한 책임감을 보존하면서 달성될 수 있다.

주요 한계로는 맥락적 뉘앙스 포착의 도전, 특정 결정 유형에 대한 인간 판단의 환원 불가능성, 잠재적 기술 의존에 대한 우려가 있다. 이러한 한계는 GPTs 도구가 인간 리더십 역량의 대체물이 아닌 보완물로 이해되어야 함을 시사한다.

AI 역량이 계속 발전하고 조직 워크플로우로의 통합이 심화됨에 따라, 이러한 기술이 인간 리더십과 어떻게 상호작용하는지 이해하는 것이 점점 더 중요해질 것이다. 본 연구는 그러한 이해의 기초를 제공하고, 인간과 인공 지능이 함께 협력하여 교육 및 기업 맥락에서 리더십 의사결정의 품질을 향상시키는 미래를 향해 방향을 제시한다."""},
]

# 참고문헌
references = [
    "Alase, A. O. (2017). The Interpretative Phenomenological Analysis (IPA): A Guide to a Good Qualitative Research Approach. International Journal of Education and Literacy Studies, 5(2), 9-19.",
    "Bass, B. M. (1985). Leadership and Performance Beyond Expectations. Free Press.",
    "Braun, V., & Clarke, V. (2006). Using Thematic Analysis in Psychology. Qualitative Research in Psychology, 3(2), 77-101.",
    "Burns, J. M. (1978). Leadership. Harper & Row.",
    "Collins, J. (2001). Good to Great: Why Some Companies Make the Leap and Others Don't. HarperBusiness.",
    "Creswell, J. W., & Creswell, J. D. (2018). Research Design: Qualitative, Quantitative, and Mixed Methods Approaches (5th ed.). Sage.",
    "Davenport, T. H., & Kirby, J. (2016). Only Humans Need Apply: Winners and Losers in the Age of Smart Machines. Harper Business.",
    "Davis, F. D. (1989). Perceived Usefulness, Perceived Ease of Use, and User Acceptance of Information Technology. MIS Quarterly, 13(3), 319-340.",
    "Eloundou, T., Manning, S., Mishkin, P., & Rock, D. (2024). GPTs are GPTs: An Early Look at the Labor Market Impact Potential of Large Language Models. Science, 384(6702), eadj0998.",
    "Engeler, I., Holm, A., Blom, J. K. R., Pedersen, M. U., & Jakobsen, A. (2024). The Unequal Adoption of ChatGPT Exacerbates Existing Inequalities Among Workers. Proceedings of the National Academy of Sciences, 121(52), e2414972121.",
    "Greenleaf, R. K. (1977). Servant Leadership: A Journey into the Nature of Legitimate Power and Greatness. Paulist Press.",
    "Kahneman, D. (2011). Thinking, Fast and Slow. Farrar, Straus and Giroux.",
    "Kotter, J. P. (1996). Leading Change. Harvard Business Review Press.",
    "Kouzes, J. M., & Posner, B. Z. (2017). The Leadership Challenge: How to Make Extraordinary Things Happen in Organizations (6th ed.). Jossey-Bass.",
    "Lincoln, Y. S., & Guba, E. G. (1985). Naturalistic Inquiry. Sage.",
    "Mintzberg, H. (1975). The Manager's Job: Folklore and Fact. Harvard Business Review, 53(4), 49-61.",
    "Noy, S., & Zhang, W. (2023). Experimental Evidence on the Productivity Effects of Generative Artificial Intelligence. Science, 381(6654), 187-192.",
    "OpenAI. (2023, November 6). Introducing GPTs. OpenAI Blog. https://openai.com/index/introducing-gpts/",
    "OpenAI. (2024, January 10). Introducing the GPT Store. OpenAI Blog. https://openai.com/index/introducing-the-gpt-store/",
    "Robinson, C., & Williams, H. (2024). Interpretative Phenomenological Analysis: Learnings from Employing IPA as a Qualitative Methodology in Educational Research. The Qualitative Report, 29(4), 939-952.",
    "Senge, P. M. (1990). The Fifth Discipline: The Art and Practice of the Learning Organization. Doubleday.",
    "Simon, H. A. (1987). Making Management Decisions: The Role of Intuition and Emotion. Academy of Management Perspectives, 1(1), 57-64.",
    "Smith, J. A., Flowers, P., & Larkin, M. (2009). Interpretative Phenomenological Analysis: Theory, Method and Research. Sage.",
    "Van Manen, M. (1990). Researching Lived Experience: Human Science for an Action Sensitive Pedagogy. SUNY Press.",
    "Venkatesh, V., Morris, M. G., Davis, G. B., & Davis, F. D. (2003). User Acceptance of Information Technology: Toward a Unified View. MIS Quarterly, 27(3), 425-478.",
]

def create_shinsa_docx():
    """신학과사회 2025년 형식 DOCX 생성"""

    doc = Document()

    # ===== 페이지 설정 (신국판 152x225mm) =====
    section = doc.sections[0]
    section.page_width = Mm(152)
    section.page_height = Mm(225)
    section.top_margin = Mm(24)
    section.bottom_margin = Mm(25)
    section.left_margin = Mm(25)
    section.right_margin = Mm(23)

    # ===== 헤더 =====
    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_run = header_para.add_run(f"「신학과 사회」 {paper_data['volume']}({paper_data['issue']}) {paper_data['year']}")
    set_korean_font(header_run, "바탕", 8.1)

    page_range = header.add_paragraph()
    page_range_run = page_range.add_run(f"pp. {paper_data['start_page']} - {paper_data['end_page']}")
    set_korean_font(page_range_run, "바탕", 8.1)
    page_range.paragraph_format.space_after = Pt(15)

    # ===== 제목 (14pt, 굵게, 가운데) =====
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(paper_data["title"])
    set_korean_font(title_run, "바탕", 14)
    title_run.bold = True
    title_para.paragraph_format.space_before = Pt(20)
    title_para.paragraph_format.space_after = Pt(8)

    # ===== 부제 (12.3pt) =====
    if paper_data.get("subtitle"):
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.add_run(paper_data["subtitle"])
        set_korean_font(subtitle_run, "바탕", 12.3)
        subtitle_para.paragraph_format.space_after = Pt(15)

    # ===== 저자 (11pt, 띄어쓰기) =====
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_spaced = " ".join(paper_data["author"])
    author_run = author_para.add_run(author_spaced)
    set_korean_font(author_run, "바탕", 11)
    author_run.bold = True

    # 저자 각주 기호
    symbol_run = author_para.add_run("*")
    symbol_run.font.superscript = True
    set_korean_font(symbol_run, "바탕", 8)
    author_para.paragraph_format.space_after = Pt(20)

    # ===== 구분선 =====
    doc.add_paragraph("─" * 35)

    # ===== 국문초록 =====
    abstract_title = doc.add_paragraph()
    abstract_title_run = abstract_title.add_run("국문초록")
    set_korean_font(abstract_title_run, "바탕", 9)
    abstract_title_run.bold = True
    abstract_title.paragraph_format.space_after = Pt(8)

    abstract_para = doc.add_paragraph()
    abstract_run = abstract_para.add_run(paper_data["abstract_kr"])
    set_korean_font(abstract_run, "바탕", 8.5)
    add_paragraph_spacing(abstract_para, 1.5, 8)
    abstract_para.paragraph_format.left_indent = Mm(10)
    abstract_para.paragraph_format.right_indent = Mm(10)

    # 주제어
    keywords_para = doc.add_paragraph()
    keywords_label = keywords_para.add_run("주제어: ")
    set_korean_font(keywords_label, "바탕", 8.5)
    keywords_label.bold = True
    keywords_run = keywords_para.add_run(", ".join(paper_data["keywords_kr"]))
    set_korean_font(keywords_run, "바탕", 8.5)
    keywords_para.paragraph_format.left_indent = Mm(10)

    # ===== 구분선 =====
    doc.add_paragraph("─" * 35)

    # ===== 본문 =====
    for sec in sections:
        if sec["title"]:
            sec_para = doc.add_paragraph()
            sec_run = sec_para.add_run(f"{sec['number']} {sec['title']}")

            if sec["level"] == 1:
                set_korean_font(sec_run, "바탕", 13)
                sec_run.bold = True
                sec_para.paragraph_format.space_before = Pt(18)
            elif sec["level"] == 2:
                set_korean_font(sec_run, "바탕", 11)
                sec_run.bold = True
                sec_para.paragraph_format.space_before = Pt(12)

            sec_para.paragraph_format.space_after = Pt(6)

        if sec["content"]:
            for para_text in sec["content"].split("\n\n"):
                if para_text.strip():
                    body_para = doc.add_paragraph()
                    body_run = body_para.add_run(para_text.strip())
                    set_korean_font(body_run, "바탕", 10.3)
                    add_paragraph_spacing(body_para, 1.6, 6)
                    body_para.paragraph_format.first_line_indent = Mm(5)

    # ===== 구분선 =====
    doc.add_paragraph("─" * 35)

    # ===== 참고문헌 =====
    ref_title = doc.add_paragraph()
    ref_title_run = ref_title.add_run("참고문헌")
    set_korean_font(ref_title_run, "바탕", 13)
    ref_title_run.bold = True
    ref_title.paragraph_format.space_before = Pt(18)
    ref_title.paragraph_format.space_after = Pt(10)

    for ref in references:
        ref_para = doc.add_paragraph()
        ref_run = ref_para.add_run(ref)
        # 영문 참고문헌은 Times New Roman
        if ref[0].isascii():
            ref_run.font.name = "Times New Roman"
            ref_run.font.size = Pt(10.3)
        else:
            set_korean_font(ref_run, "바탕", 10.3)
        ref_para.paragraph_format.left_indent = Mm(5)
        ref_para.paragraph_format.first_line_indent = Mm(-5)
        ref_para.paragraph_format.space_after = Pt(3)

    # ===== 구분선 =====
    doc.add_paragraph("─" * 35)

    # ===== 영문 초록 =====
    abstract_en_title = doc.add_paragraph()
    abstract_en_run = abstract_en_title.add_run("Abstract")
    abstract_en_run.font.name = "Times New Roman"
    abstract_en_run.font.size = Pt(9)
    abstract_en_run.bold = True

    abstract_en_para = doc.add_paragraph()
    abstract_en_content = abstract_en_para.add_run(paper_data["abstract_en"])
    abstract_en_content.font.name = "Times New Roman"
    abstract_en_content.font.size = Pt(8.5)
    add_paragraph_spacing(abstract_en_para, 1.5, 8)
    abstract_en_para.paragraph_format.left_indent = Mm(10)
    abstract_en_para.paragraph_format.right_indent = Mm(10)

    # Keywords
    keywords_en_para = doc.add_paragraph()
    keywords_en_label = keywords_en_para.add_run("Keywords: ")
    keywords_en_label.font.name = "Times New Roman"
    keywords_en_label.font.size = Pt(8.5)
    keywords_en_label.bold = True
    keywords_en_content = keywords_en_para.add_run(", ".join(paper_data["keywords_en"]))
    keywords_en_content.font.name = "Times New Roman"
    keywords_en_content.font.size = Pt(8.5)
    keywords_en_para.paragraph_format.left_indent = Mm(10)

    # ===== 구분선 =====
    doc.add_paragraph("─" * 35)

    # ===== 저자 각주 =====
    footnote_para = doc.add_paragraph()
    footnote_run = footnote_para.add_run(f"* {paper_data['affiliation']}/ {paper_data['field']}")
    set_korean_font(footnote_run, "바탕", 8.5)

    # ===== 저장 =====
    output_path = Path(r"C:\Users\sshin\OneDrive - Global Education Research Institutute\바탕 화면\GPTs_Leadership_신사형식_2025.docx")
    doc.save(str(output_path))

    print(f"✓ 신학과사회 2025년 형식 DOCX 생성 완료!")
    print(f"  경로: {output_path}")
    print(f"\n적용된 형식:")
    print(f"  - 페이지: 신국판 (152x225mm)")
    print(f"  - 마진: 상24/하25/좌25/우23mm")
    print(f"  - 폰트: 바탕체")
    print(f"  - 제목: 14pt, 부제: 12.3pt")
    print(f"  - 본문: 10.3pt, 줄간격: 160%")
    print(f"  - 초록: 8.5pt, 섹션제목: 13pt")

if __name__ == "__main__":
    create_shinsa_docx()
